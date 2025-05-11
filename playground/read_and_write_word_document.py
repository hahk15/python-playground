import docx

# Path to document (end with .docx)
PATH_TO_DOCX_DOCUMENT = "playground/data/python-library-reference.docx"

# Read content by specific path
content = docx.Document(docx=PATH_TO_DOCX_DOCUMENT)

# Add sample output
content.add_heading("This is another heading")
content.add_page_break()
content.add_heading("Part 2", level=2)
content.add_table(1, 2)

# Save to another file
content.save("playground/data/python-library-reference-update.docx")
